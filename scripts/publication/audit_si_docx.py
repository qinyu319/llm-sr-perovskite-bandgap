import json
import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from lxml import etree


ROOT = Path(__file__).resolve().parents[2]
DOCX = Path(sys.argv[1]) if len(sys.argv) > 1 else ROOT / "paper" / "si_REVISED.docx"
NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def fail(message):
    failures.append(message)


doc = Document(DOCX)
paragraph_text = "\n".join(p.text for p in doc.paragraphs)
table_text = "\n".join(
    cell.text for table in doc.tables for row in table.rows for cell in row.cells
)
text = f"{paragraph_text}\n{table_text}"
headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
failures = []

required_headings = [f"Supplementary Note S{i}." for i in range(1, 16)]
for heading in required_headings:
    if not any(h.startswith(heading) for h in headings):
        fail(f"Missing heading: {heading}")

required_phrases = [
    "Eg = 1.55527 - 1.10253 Sn + 0.34320 Br + 1.61932 Cl + 0.12268 Cs",
    "0.91702 Sn2 + 0.36607 Br2 - 0.22716 Cs·Sn - 0.32528 Cs·Cl",
    "The project archive does not contain a verified external panel",
    "Exact prompt template for M0",
    "Exact prompt template for M1",
    "Exact prompt template for M2",
    "Exact prompt template for M3",
    "Exact prompt template for M4",
    "Immutable 518-row training workbook",
    "Immutable 92-row held-out workbook",
]
for phrase in required_phrases:
    if phrase not in text:
        fail(f"Missing required phrase: {phrase}")

if len(doc.inline_shapes) != 13:
    fail(f"Expected 13 inline figures, found {len(doc.inline_shapes)}")

table_geometry = []
with zipfile.ZipFile(DOCX) as archive:
    xml = etree.fromstring(archive.read("word/document.xml"))
    tables = xml.xpath("//w:tbl", namespaces=NS)
    for index, table in enumerate(tables, start=1):
        width_nodes = table.xpath("./w:tblPr/w:tblW", namespaces=NS)
        grid_cols = table.xpath("./w:tblGrid/w:gridCol", namespaces=NS)
        width = int(width_nodes[0].get(f"{{{NS['w']}}}w")) if width_nodes else None
        grid_sum = sum(int(node.get(f"{{{NS['w']}}}w")) for node in grid_cols)
        table_geometry.append({"table": index, "width": width, "grid_sum": grid_sum})
        if width != 9360:
            fail(f"Table {index} width is {width}, expected 9360")
        if grid_sum != 9360:
            fail(f"Table {index} grid sum is {grid_sum}, expected 9360")

caption_numbers = sorted(
    {int(match) for match in re.findall(r"Figure S(\d+)\.", text)}
)
if caption_numbers != list(range(1, 14)):
    fail(f"Figure captions are incomplete: {caption_numbers}")

table_caption_numbers = sorted(
    {int(match) for match in re.findall(r"Table S(\d+)[A-B]?\.", text)}
)
if table_caption_numbers != list(range(0, 28)):
    fail(f"Table captions are incomplete: {table_caption_numbers}")

report = {
    "path": str(DOCX.resolve()),
    "size_bytes": DOCX.stat().st_size,
    "paragraphs": len(doc.paragraphs),
    "headings": len(headings),
    "tables": len(doc.tables),
    "inline_figures": len(doc.inline_shapes),
    "figure_captions": caption_numbers,
    "table_captions": table_caption_numbers,
    "failures": failures,
}
print(json.dumps(report, indent=2, ensure_ascii=False))
sys.exit(1 if failures else 0)
