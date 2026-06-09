from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
BUILD = ROOT / "archive" / "publication_build" / "main"
FIG = ROOT / "figures" / "main"
OUTPUT = (
    ROOT
    / "reproduced"
    / "manuscript"
    / "Reproducible_Physics_Constrained_LLM_SR_Perovskite_Bandgap_review.docx"
)
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

NAVY = RGBColor(31, 78, 121)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(89, 89, 89)
LIGHT_GRAY = "F2F2F2"
MID_GRAY = "D9E2F3"
FONT = "Times New Roman"


def set_run_font(run, size=11, bold=None, italic=None, color=BLACK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "B7B7B7")


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
    set_table_borders(table)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=9, color=GRAY)


def add_custom_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
    ]
    existing_num = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs, default=0) + 1
    num_id = max(existing_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, ind])
    lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId")
    abs_id.set(qn("w:val"), str(abstract_id))
    num.append(abs_id)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def add_inline(paragraph, text, size=11, color=BLACK):
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            set_run_font(paragraph.add_run(text[cursor : match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), size=size, bold=True, color=color)
        elif token.startswith("*"):
            set_run_font(paragraph.add_run(token[1:-1]), size=size, italic=True, color=color)
        else:
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Courier New")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Courier New")
            run.font.size = Pt(size - 1)
            run.font.color.rgb = color
        cursor = match.end()
    if cursor < len(text):
        set_run_font(paragraph.add_run(text[cursor:]), size=size, color=color)


def add_equation(doc):
    p = doc.add_paragraph(style="Equation")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parts = [
        ("E", {"italic": True}),
        ("g", {"subscript": True}),
        (" = 1.55527 - 1.10253 Sn + 0.34320 Br + 1.61932 Cl + 0.12268 Cs", {}),
        (" + 0.91702 Sn", {}),
        ("2", {"superscript": True}),
        (" + 0.36607 Br", {}),
        ("2", {"superscript": True}),
        (" - 0.22716 CsSn - 0.32528 CsCl", {}),
    ]
    for text, opts in parts:
        run = p.add_run(text)
        set_run_font(run, size=10.5)
        run.font.italic = opts.get("italic", False)
        run.font.subscript = opts.get("subscript", False)
        run.font.superscript = opts.get("superscript", False)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline(p, text, size=9, color=GRAY)
    return p


def add_figure(doc, filename, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(FIG / filename), width=Inches(6.25))
    for drawing in run._r.findall(qn("w:drawing")):
        doc_pr = drawing.find(".//wp:docPr", {"wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"})
        if doc_pr is not None:
            doc_pr.set("descr", caption[:250])
    add_caption(doc, caption)


def add_table(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        set_cell_shading(hdr.cells[i], MID_GRAY)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(str(header)), size=font_size, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i in (0, len(headers) - 1) else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_run_font(p.add_run(str(value)), size=font_size)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    return table


def table_dataset(doc):
    add_caption(doc, "Table 1. Dataset and fixed-split summary.")
    headers = ["Cohort", "n", "Eg mean ± SD (eV)", "Median (eV)", "1.5-1.7 eV", "Eg >= 2.2 eV", "Cl > 0"]
    rows = [
        ["Training", "518", "1.651 ± 0.310", "1.606", "268 (51.7%)", "36 (6.9%)", "27"],
        ["Held-out test", "92", "1.695 ± 0.398", "1.615", "43 (46.7%)", "10 (10.9%)", "8"],
        ["Strict-QC training", "507", "1.653 ± 0.312", "1.604", "Not rebalanced", "Not rebalanced", "QC subset"],
    ]
    add_table(doc, headers, rows, [1250, 650, 1600, 1100, 1400, 1500, 1760], font_size=8.1)


def table_coefficients(doc):
    add_caption(doc, "Table 2. Final M4 coefficient diagnostics.")
    coef = pd.read_csv(ROOT / "final_m4_diagnostics" / "final_M4_coefficients_CI_bootstrap.csv")
    coef = coef[coef["Raw term"].astype(str) != "Intercept"]
    rows = []
    for _, r in coef.iterrows():
        rows.append(
            [
                str(r["Term"]),
                f'{float(r["Coefficient"]):.5f}',
                f'[{float(r["OLS CI low"]):.3f}, {float(r["OLS CI high"]):.3f}]',
                f'[{float(r["Bootstrap CI low"]):.3f}, {float(r["Bootstrap CI high"]):.3f}]',
                f'{float(r["Bootstrap sign stability"]):.3f}',
                f'{float(r["VIF"]):.2f}',
            ]
        )
    add_table(
        doc,
        ["Term", "Coefficient", "OLS 95% CI", "Bootstrap 95% CI", "Sign stability", "VIF"],
        rows,
        [1050, 1250, 1800, 2050, 1650, 1560],
        font_size=8.3,
    )


def table_benchmarks(doc):
    add_caption(doc, "Table 3. Main predictive and symbolic-regression benchmarks.")
    rows = [
        ["Final M4", "8-term OLS formula", "0.0575", "0.0606", "Frozen; high coefficient stability"],
        ["Gaussian process", "Matérn-1.5", "0.0513", "0.0486", "Best predictive benchmark"],
        ["GBRT", "Tree ensemble", "0.0532", "0.0520", "SHAP-interpretable black box"],
        ["Random forest", "Tree ensemble", "0.0597", "0.0559", "Higher model complexity"],
        ["XGBoost", "Boosted trees", "0.0513", "0.0606", "Best CV among tree models"],
        ["Exhaustive-6", "Anion polynomial subset", "0.0629", "0.0667", "Deterministic; excludes A-site"],
        ["PySR-P", "Five symbolic runs", "n/a", "0.0683 ± 0.0096", "Five unique selected structures"],
        ["GPLearn", "Compact GP baseline", "0.1453", "0.1773 ± 0.0228", "Exploratory limited-budget run"],
    ]
    add_table(
        doc,
        ["Method", "Model/search type", "CV RMSE (eV)", "Test RMSE (eV)", "Interpretation"],
        rows,
        [1350, 1800, 1250, 1350, 3610],
        font_size=7.8,
    )


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.85)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(10.5)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.widow_control = True

for style_name, size, before, after, color in [
    ("Heading 1", 14, 14, 6, NAVY),
    ("Heading 2", 12, 11, 4, NAVY),
    ("Heading 3", 11, 8, 3, NAVY),
]:
    st = styles[style_name]
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
    st.paragraph_format.keep_together = True

if "Equation" not in styles:
    eq_style = styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
else:
    eq_style = styles["Equation"]
eq_style.paragraph_format.space_before = Pt(6)
eq_style.paragraph_format.space_after = Pt(6)
eq_style.paragraph_format.keep_together = True

caption_style = styles["Caption"]
caption_style.font.name = FONT
caption_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
caption_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
caption_style.font.size = Pt(9)
caption_style.font.italic = False
caption_style.font.color.rgb = GRAY
caption_style.paragraph_format.space_before = Pt(2)
caption_style.paragraph_format.space_after = Pt(7)
caption_style.paragraph_format.line_spacing = 1.0

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hp.paragraph_format.space_after = Pt(0)
set_run_font(hp.add_run("Research Article | Physics-Constrained LLM-SR"), size=8.5, color=GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run_font(fp.add_run("Page "), size=9, color=GRAY)
add_page_field(fp)

num_id = add_custom_numbering(doc)

lines = (BUILD / "manuscript.md").read_text(encoding="utf-8").splitlines()
paragraph_buffer = []
title_done = False


def flush_paragraph():
    global paragraph_buffer
    if not paragraph_buffer:
        return
    text = " ".join(x.strip() for x in paragraph_buffer).strip()
    p = doc.add_paragraph()
    add_inline(p, text, size=10.5)
    paragraph_buffer = []


for raw in lines:
    line = raw.rstrip()
    if not line.strip():
        flush_paragraph()
        continue
    if line.startswith("[[FIGURE:"):
        flush_paragraph()
        inner = line[len("[[FIGURE:") : -2]
        filename, caption = inner.split("|", 1)
        add_figure(doc, filename, caption)
        continue
    if line.startswith("[[EQUATION:"):
        flush_paragraph()
        add_equation(doc)
        continue
    if line == "[[TABLE:dataset]]":
        flush_paragraph()
        table_dataset(doc)
        continue
    if line == "[[TABLE:coefficients]]":
        flush_paragraph()
        table_coefficients(doc)
        continue
    if line == "[[TABLE:benchmarks]]":
        flush_paragraph()
        table_benchmarks(doc)
        continue
    if line.startswith("# "):
        flush_paragraph()
        text = line[2:].strip()
        if not title_done:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(9)
            p.paragraph_format.keep_with_next = True
            set_run_font(p.add_run(text), size=19, bold=True, color=NAVY)
            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub.paragraph_format.space_after = Pt(14)
            set_run_font(sub.add_run("Research Article"), size=10, italic=True, color=GRAY)
            title_done = True
        else:
            p = doc.add_paragraph(text, style="Heading 1")
            for run in p.runs:
                set_run_font(run, size=14, bold=True, color=NAVY)
        continue
    if line.startswith("## "):
        flush_paragraph()
        text = line[3:].strip()
        p = doc.add_paragraph(text, style="Heading 2")
        for run in p.runs:
            set_run_font(run, size=12, bold=True, color=NAVY)
        continue
    if line.startswith("### "):
        flush_paragraph()
        text = line[4:].strip()
        p = doc.add_paragraph(text, style="Heading 3")
        for run in p.runs:
            set_run_font(run, size=11, bold=True, color=NAVY)
        continue
    if re.match(r"^\d+\.\s", line):
        flush_paragraph()
        text = re.sub(r"^\d+\.\s*", "", line)
        p = doc.add_paragraph()
        apply_num(p, num_id)
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        add_inline(p, text, size=9.2)
        continue
    paragraph_buffer.append(line)

flush_paragraph()

# Prevent the first abstract heading from feeling detached from the title.
for p in doc.paragraphs:
    if p.text == "Abstract":
        p.paragraph_format.space_before = Pt(4)
        break

doc.core_properties.title = "Reproducible Physics-Constrained LLM-Assisted Symbolic Regression for Interpretable Band-Gap Modeling in Hybrid Perovskites"
doc.core_properties.subject = "Hybrid perovskite band-gap modeling"
doc.core_properties.author = ""
doc.core_properties.keywords = "perovskites, symbolic regression, interpretable machine learning, LLM-SR"
doc.core_properties.comments = "Generated from the audited project results."

doc.save(OUTPUT)
print(OUTPUT)
