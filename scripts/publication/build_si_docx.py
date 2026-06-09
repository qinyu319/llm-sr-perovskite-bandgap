from __future__ import annotations

import json
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
BUILD = ROOT / "archive" / "publication_build" / "si"
FIG = ROOT / "figures" / "supplementary"
DATA = BUILD / "data"
OUTPUT = (
    ROOT
    / "reproduced"
    / "si"
    / "Supplementary_Information_Physics_Constrained_LLM_SR_review.docx"
)
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

FONT = "Times New Roman"
MONO = "Courier New"
NAVY = RGBColor(31, 78, 121)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(89, 89, 89)
RED = RGBColor(176, 58, 46)
GREEN = RGBColor(46, 125, 50)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
CAUTION_FILL = "FFF2CC"
RED_FILL = "FCE8E6"


def set_run_font(run, size=10.25, bold=None, italic=None, color=BLACK, name=FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    rpr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_borders(table, color="B7B7B7", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for cell, width in zip(row.cells, widths_dxa):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
    set_table_borders(table)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run_font(run, size=8.5, color=GRAY)


def shade_paragraph(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color="B7B7B7", size="6"):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def add_inline(paragraph, text, size=10.25, color=BLACK):
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            set_run_font(paragraph.add_run(text[cursor : match.start()]), size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            set_run_font(paragraph.add_run(token[2:-2]), size=size, bold=True, color=color)
        elif token.startswith("*"):
            set_run_font(paragraph.add_run(token[1:-1]), size=size, italic=True, color=color)
        else:
            set_run_font(paragraph.add_run(token[1:-1]), size=size - 0.7, color=color, name=MONO)
        cursor = match.end()
    if cursor < len(text):
        set_run_font(paragraph.add_run(text[cursor:]), size=size, color=color)


def add_body(text, size=10.25, after=5, keep=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_together = keep
    add_inline(p, text, size=size)
    return p


def add_heading(text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    for run in p.runs:
        set_run_font(
            run,
            size={1: 14, 2: 12, 3: 10.75}[level],
            bold=True,
            color=NAVY,
        )
    return p


def add_caption(text):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_inline(p, text, size=8.5, color=GRAY)
    return p


def add_callout(label, text, fill=LIGHT_FILL, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.10)
    p.paragraph_format.right_indent = Inches(0.10)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.1
    shade_paragraph(p, fill)
    set_run_font(p.add_run(f"{label}: "), size=9.5, bold=True, color=color)
    add_inline(p, text, size=9.5, color=color)
    return p


def add_equation():
    p = doc.add_paragraph(style="Equation")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    parts = [
        ("E", {"italic": True}),
        ("g", {"subscript": True}),
        (" = 1.55527 - 1.10253 Sn + 0.34320 Br + 1.61932 Cl + 0.12268 Cs", {}),
        (" + 0.91702 Sn", {}),
        ("2", {"superscript": True}),
        (" + 0.36607 Br", {}),
        ("2", {"superscript": True}),
        (" - 0.22716 CsSn - 0.32528 CsCl", {}),
    ]
    for text, opts in parts:
        run = p.add_run(text)
        set_run_font(run, size=10.3)
        run.font.italic = opts.get("italic", False)
        run.font.subscript = opts.get("subscript", False)
        run.font.superscript = opts.get("superscript", False)
    return p


def add_figure(filename, caption, width=6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(FIG / filename), width=Inches(width))
    for drawing in run._r.findall(qn("w:drawing")):
        doc_pr = drawing.find(
            ".//wp:docPr",
            {"wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"},
        )
        if doc_pr is not None:
            doc_pr.set("descr", caption[:250])
    add_caption(caption)


def add_table(caption, headers, rows, widths, font_size=8.0, alignments=None):
    add_caption(caption)
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, header in enumerate(headers):
        set_cell_shading(hdr.cells[i], HEADER_FILL)
        p = hdr.cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        set_run_font(p.add_run(str(header)), size=font_size, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            alignment = (
                alignments[i]
                if alignments
                else (WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)
            )
            p.alignment = alignment
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_run_font(p.add_run(str(value)), size=font_size)
    set_table_geometry(table, widths)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)
    return table


def add_prompt_box(number, stage, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    set_run_font(
        p.add_run(f"Box S{number}. Exact prompt template for {stage}."),
        size=9,
        bold=True,
        color=NAVY,
    )
    box = doc.add_paragraph()
    box.paragraph_format.left_indent = Inches(0.10)
    box.paragraph_format.right_indent = Inches(0.10)
    box.paragraph_format.space_before = Pt(0)
    box.paragraph_format.space_after = Pt(7)
    box.paragraph_format.line_spacing = 1.0
    box.paragraph_format.keep_together = True
    shade_paragraph(box, LIGHT_FILL)
    set_run_font(box.add_run(text.strip()), size=7.4, name=MONO)


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


summary = json.loads((DATA / "si_summary.json").read_text(encoding="utf-8"))
duplicates = pd.read_csv(DATA / "duplicate_composition_audit.csv")
split_stats = pd.read_csv(DATA / "train_test_summary.csv")
closure = pd.read_csv(DATA / "closure_audit.csv")
design = pd.read_csv(DATA / "design_candidates.csv")
coeff = pd.read_csv(ROOT / "final_m4_diagnostics" / "final_M4_coefficients_CI_bootstrap.csv")
cl2 = pd.read_csv(ROOT / "final_m4_diagnostics" / "Cl2_retention_diagnostic.csv")
ma = pd.read_csv(ROOT / "ma_asite_sensitivity" / "ma_asite_model_comparison.csv")
regularization = pd.read_csv(
    ROOT / "baselines/regularization_aic_bic" / "regularization_summary.csv"
)
external_llm = pd.read_csv(
    ROOT / "llm_repeated" / "external_api_runs" / "external_api_summary_with_median_iqr.csv"
)
group_summary = pd.read_csv(
    ROOT
    / "group_aware_sensitivity"
    / "outputs"
    / "summary_tables"
    / "group_aware_summary.csv"
)
group_heldout = pd.read_csv(
    ROOT
    / "group_aware_sensitivity"
    / "outputs"
    / "summary_tables"
    / "group_aware_heldout_group_summary.csv"
)
group_terms = pd.read_csv(
    ROOT
    / "group_aware_sensitivity"
    / "outputs"
    / "summary_tables"
    / "group_aware_term_frequency.csv"
)
pysr = json.loads((ROOT / "baselines/pysr" / "pysr_summary.json").read_text(encoding="utf-8"))
gplearn = json.loads(
    (ROOT / "baselines" / "gplearn" / "gplearn_summary.json").read_text(
        encoding="utf-8"
    )
)
repeated_summary = json.loads(
    (
        ROOT
        / "llm_repeated"
        / "repeated_runs_30"
        / "summary"
        / "experiment_summary.json"
    ).read_text(encoding="utf-8")
)
direct_summary = json.loads(
    (
        ROOT
        / "llm_repeated"
        / "raw_outputs"
        / "candidates_30_M4_codex_summary.json"
    ).read_text(encoding="utf-8")
)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(10.25)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.widow_control = True

for style_name, size, before, after, color in [
    ("Heading 1", 14, 16, 8, NAVY),
    ("Heading 2", 12, 12, 6, NAVY),
    ("Heading 3", 10.75, 8, 4, NAVY),
]:
    st = styles[style_name]
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
    st.paragraph_format.keep_together = True

if "Equation" not in styles:
    equation_style = styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
else:
    equation_style = styles["Equation"]
equation_style.paragraph_format.space_before = Pt(6)
equation_style.paragraph_format.space_after = Pt(6)
equation_style.paragraph_format.keep_together = True

caption_style = styles["Caption"]
caption_style.font.name = FONT
caption_style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
caption_style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
caption_style.font.size = Pt(8.5)
caption_style.font.color.rgb = GRAY
caption_style.font.italic = False
caption_style.paragraph_format.space_before = Pt(3)
caption_style.paragraph_format.space_after = Pt(6)
caption_style.paragraph_format.line_spacing = 1.0
caption_style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.paragraph_format.space_after = Pt(0)
set_run_font(hp.add_run("Supplementary Information | Physics-Constrained LLM-SR"), size=8, color=GRAY)
set_paragraph_border(hp, color="D9D9D9", size="4")

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run_font(fp.add_run("Page "), size=8.5, color=GRAY)
add_page_field(fp)

# Cover.
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(32)
p.paragraph_format.space_after = Pt(10)
set_run_font(p.add_run("SUPPLEMENTARY INFORMATION"), size=13, bold=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)
set_run_font(
    p.add_run(
        "Reproducible Physics-Constrained LLM-Assisted Symbolic Regression "
        "for Interpretable Band-Gap Modeling in Hybrid Perovskites"
    ),
    size=19,
    bold=True,
    color=BLACK,
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
set_run_font(p.add_run("Evidence package for the revised manuscript"), size=10.5, italic=True, color=GRAY)

add_callout(
    "Scope",
    "This Supplementary Information documents dataset integrity, closure-aware encoding, immutable split files, exact M0-M4 prompts, stage-wise selection, final-model diagnostics, non-LLM controls, LLM stochasticity, SHAP cross-checks, group-aware sensitivity, design demonstrations, and the local reproducibility manifest.",
    fill=LIGHT_FILL,
)
add_callout(
    "Version lock",
    "Throughout this document, **Final frozen 8-term M4** refers exclusively to the CsCl-containing formula below. The historical 13-term M4, the repeated-run CsBr candidate, and anion-only polynomial models are not called the final model.",
    fill=CAUTION_FILL,
)
add_equation()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
set_run_font(
    p.add_run("CV RMSE = 0.057532 eV | Test RMSE = 0.060613 eV | Test R² = 0.976545"),
    size=10,
    bold=True,
    color=NAVY,
)
add_callout(
    "Evidence boundary",
    "The archive contains a same-source held-out test set but no verified independent literature validation set. No external-validation table, parity plot, repository URL, or Zenodo DOI is fabricated in this document.",
    fill=RED_FILL,
    color=RED,
)

page_break()
add_heading("Contents", 1)
contents = [
    ("S1", "Dataset curation and duplicate-record handling"),
    ("S2", "Composition encoding and closure constraints"),
    ("S3", "Train/test split reproducibility and fairness"),
    ("S4", "Exact prompt templates and physical constraints"),
    ("S5", "Stage-wise candidate generation and M0-M4 evolution"),
    ("S6", "Final M4 freezing rule and model nomenclature"),
    ("S7", "Coefficient diagnostics, VIF, confidence intervals, and bootstrap"),
    ("S8", "MA and A-site encoding sensitivity"),
    ("S9", "Non-LLM symbolic regression and statistical controls"),
    ("S10", "Repeated LLM runs and cross-model stochasticity"),
    ("S11", "Black-box baselines and SHAP cross-check"),
    ("S12", "Group-aware splitting and applicability domain"),
    ("S13", "Independent external-validation status"),
    ("S14", "Composition-guided design demonstration"),
    ("S15", "Reproducibility package and file manifest"),
]
for label, title in contents:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(f"{label}  "), size=10, bold=True, color=NAVY)
    set_run_font(p.add_run(title), size=10)

page_break()
add_heading("Supplementary Note S1. Dataset curation and duplicate-record handling", 1)
add_body(
    "The source data were derived from the open Perovskite Database Project described by Jacobsson et al. The frozen local workbook snapshot is dated January 28, 2026. The retained scope is three-dimensional ABX3 hybrid perovskites represented by A-site fractions FA, MA, and Cs; B-site fractions Pb and Sn; X-site fractions I, Br, and Cl; and a reported optical band gap. Journal impact factor, Journal Citation Reports category, and publication-prestige filters were not used. The target distribution was not balanced, truncated, oversampled, or reshaped."
)
add_body(
    "The archive begins at a 610-row record-level worksheet. The preceding database export and row-by-row exclusion log are not available, so counts before the 610-row snapshot cannot be reconstructed. Table S1 separates directly auditable steps from this upstream evidence gap."
)
add_table(
    "Table S1. Dataset curation workflow and evidence status.",
    ["Step", "Criterion or operation", "Retained records", "Evidence status"],
    [
        ["1", "Source database export", "Not recoverable", "Pre-610 export not archived"],
        ["2", "Restrict to 3D ABX3; A={FA, MA, Cs}; B={Pb, Sn}; X={I, Br, Cl}; Eg present", "610", "Documented scope; upstream row log absent"],
        ["3", "No IF/JCR or publication-prestige filter", "610", "Documented policy"],
        ["4", "No artificial Eg balancing or truncation", "610", "Distribution retained"],
        ["5", "Retain record-level duplicate measurements", "610", "Primary analysis policy"],
        ["6", "Immutable 85/15 workbooks", "518 / 92", "Checksummed train/test inputs"],
        ["7", "Strict closure sensitivity without renormalization", "507 training rows", "Group-aware analysis only"],
    ],
    [650, 4650, 1450, 2610],
    font_size=7.8,
)
add_figure(
    "figure_s1_curation_flow.png",
    "Figure S1. Evidence-aware data workflow. The archived record-level snapshot contains 610 rows and is exactly partitioned into the 518-row training and 92-row test workbooks. The upstream extraction log is not present. Strict closure filtering was used only for the group-aware sensitivity workflow.",
)
add_body(
    f"The 610-row worksheet contains {summary['dataset']['unique_full_rows']} unique complete rows and {summary['dataset']['unique_compositions']} unique compositions when Eg is excluded. There are {summary['dataset']['duplicate_composition_groups']} composition groups with more than one measurement, corresponding to {summary['dataset']['duplicate_composition_excess']} records beyond one record per composition. Eleven records are exact duplicates of an existing complete row, including Eg. The primary analysis retained these measurements as separate records because the archive does not contain measurement-condition metadata sufficient to justify averaging. The auxiliary `processed` worksheet contains 519 means derived from the 575-row Cl-free subset and was not used to construct Final M4."
)
duplicate_rows = [
    [
        row["composition"],
        int(row["count"]),
        f'{row["min"]:.3f}',
        f'{row["max"]:.3f}',
        f'{row["mean"]:.3f}',
        "" if pd.isna(row["std"]) else f'{row["std"]:.3f}',
    ]
    for _, row in duplicates.iterrows()
]
add_table(
    "Table S2. Complete duplicate-composition audit. All listed measurements were retained as separate records in the primary 610-row analysis.",
    ["Composition fractions", "n", "Eg min", "Eg max", "Eg mean", "Eg SD"],
    duplicate_rows,
    [4800, 550, 900, 900, 900, 1310],
    font_size=7.0,
    alignments=[
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.CENTER,
    ],
)
add_figure(
    "figure_s2_archived_distributions.png",
    "Figure S2. Archived record-level and auxiliary composition-averaged band-gap distributions. Panel (a) is the 610-record primary dataset. Panel (b) is the 519-row composition-mean sheet derived from the Cl-free subset and is shown only to document the workbook contents.",
)
add_callout(
    "Interpretation",
    "The difference between record-level and composition-mean distributions arises from duplicate handling and the Cl-free scope of the auxiliary sheet. It is not evidence of target balancing.",
)

add_heading("Supplementary Note S2. Composition encoding and closure constraints", 1)
add_body(
    "The regression variables are site fractions rather than a larger physicochemical descriptor set. Closure makes several fractions algebraically dependent. With an intercept, including all fractions from a site would create exact linear dependence. The final representation therefore uses Sn for the B site, Br and Cl for the X site, and a restricted Cs correction for the A site."
)
add_table(
    "Table S3. Independent and dependent composition variables.",
    ["Site", "Fractions", "Independent variables in Final M4", "Closure relation and omitted baseline"],
    [
        ["B", "Pb, Sn", "Sn", "Pb = 1 - Sn"],
        ["X", "I, Br, Cl", "Br, Cl", "I = 1 - Br - Cl"],
        ["A", "FA, MA, Cs", "Cs", "FA + MA = 1 - Cs; MA sensitivity tested separately"],
    ],
    [800, 1700, 2600, 4260],
    font_size=8.2,
)
add_body(
    "Fractions were not renormalized. A strict tolerance of 10^-6 was applied in the group-aware sensitivity analysis. Eleven training rows failed at least one A-, B-, or X-site closure check and were excluded from that sensitivity workflow, leaving 507 rows. The primary final-model analysis retained the immutable 518-row training workbook so that model comparisons remained consistent with the frozen split. Two test rows had closure issues and were retained only as reference observations; the test set was never used for selection."
)
closure_rows = [
    [
        row["cohort"],
        int(row["n"]),
        int(row["closure_pass"]),
        int(row["closure_issue"]),
        f'{row["max_abs_A_error"]:.4f}',
        f'{row["max_abs_B_error"]:.4f}',
        f'{row["max_abs_X_error"]:.4f}',
    ]
    for _, row in closure.iterrows()
]
add_table(
    "Table S4. Closure audit for immutable split workbooks.",
    ["Cohort", "n", "Pass", "Issue", "Max |A-1|", "Max |B-1|", "Max |X-1|"],
    closure_rows,
    [1500, 750, 850, 850, 1800, 1800, 1810],
    font_size=8.0,
)
add_callout(
    "No sign constraints",
    "Coefficient signs were not imposed manually. Physical knowledge entered through closure-aware variables, the admissible quadratic term space, stage-wise expansion, and diagnostic filtering.",
)

add_heading("Supplementary Note S3. Train/test split reproducibility and fairness", 1)
add_body(
    "The supplied split was fixed before the present analysis and was treated as immutable. The training workbook contains 518 rows and the test workbook 92 rows. Their combined complete-row multiset is exactly equal to the 610-row archived worksheet, so no record was added or lost during splitting. The original split-generation seed is not present in the archive and was not recovered by testing scikit-learn `train_test_split(test_size=0.15)` seeds 0-199999. Reproducibility is therefore based on the checksummed split files rather than a reconstructed seed."
)
add_body(
    "All candidate generation, OLS fitting during cross-validation, pruning, bootstrap analysis, and final freezing used the training workbook only. The held-out workbook was evaluated after the Final M4 structure was locked. Fixed model comparisons used five-fold `KFold(n_splits=5, shuffle=True, random_state=42)`; repeated workflow experiments used their separately archived CV seed 20260607."
)
variables = ["Eg", "Sn", "Br", "Cl", "Cs", "MA", "FA"]
stat_rows = []
for variable in variables:
    train_row = split_stats[(split_stats.cohort == "Training") & (split_stats.variable == variable)].iloc[0]
    test_row = split_stats[(split_stats.cohort == "Held-out test") & (split_stats.variable == variable)].iloc[0]
    stat_rows.append(
        [
            variable,
            f'{train_row["mean"]:.4f} ± {train_row["sd"]:.4f}',
            f'{train_row["median"]:.4f}',
            f'{train_row["min"]:.4f}-{train_row["max"]:.4f}',
            f'{test_row["mean"]:.4f} ± {test_row["sd"]:.4f}',
            f'{test_row["median"]:.4f}',
            f'{test_row["min"]:.4f}-{test_row["max"]:.4f}',
        ]
    )
add_table(
    "Table S5. Summary statistics for immutable train and test splits.",
    ["Variable", "Train mean ± SD", "Train median", "Train range", "Test mean ± SD", "Test median", "Test range"],
    stat_rows,
    [850, 1650, 1150, 1400, 1650, 1150, 1510],
    font_size=7.6,
)
add_figure(
    "figure_s3_train_test_distributions.png",
    "Figure S3. Train/test distribution comparison for Eg and composition fractions. The dominant density lies in the I/Br-rich composition manifold; Cl-containing observations remain sparse.",
)
add_table(
    "Table S6. Duplicate and overlap audit across the fixed split.",
    ["Audit item", "Result", "Interpretation"],
    [
        ["Exact complete rows shared across train/test", f"{summary['split']['shared_exact_full_rows']} unique; {summary['split']['test_rows_exactly_repeated_from_training']} test rows", "Potential optimistic leakage for two records"],
        ["Exact composition groups shared across train/test", f"{summary['split']['shared_full_composition_groups']} groups; {summary['split']['test_rows_in_shared_full_compositions']} test rows", "Same fractions, Eg may differ"],
        ["Shared Final-M4 feature combinations", f"{summary['split']['shared_model_feature_combinations']} combinations; {summary['split']['test_rows_shared_model_features']} test rows", "Interpolation at model-feature level"],
        ["Unseen Final-M4 feature combinations", f"{summary['split']['test_rows_unseen_model_features']} test rows", "Harder local generalization"],
        ["M4 RMSE on shared feature combinations", f"{summary['split']['shared_feature_rmse']:.6f} eV", "Lower-error interpolation subset"],
        ["M4 RMSE on unseen feature combinations", f"{summary['split']['unseen_feature_rmse']:.6f} eV", "Higher-error subset"],
    ],
    [3100, 2200, 4060],
    font_size=7.8,
)
add_callout(
    "Fairness limitation",
    "Random splitting does not eliminate composition-family overlap. This is why the group-aware analyses in Note S12 are required and why the random-split test error is interpreted as interpolation performance rather than unrestricted extrapolation.",
    fill=CAUTION_FILL,
)

page_break()
add_heading("Supplementary Note S4. Exact prompt templates and physical constraints", 1)
add_body(
    "The LLM proposed symbolic structures only. It did not receive coefficient values, raw data rows, or held-out test metrics in the frozen direct templates. Numerical coefficients were estimated locally by OLS. Candidate selection was performed by local training-set cross-validation and deterministic diagnostics."
)
add_table(
    "Table S7. Global prompt and fitting constraints.",
    ["Constraint", "Setting"],
    [
        ["Variables", "Composition fractions only"],
        ["Operators", "Addition, subtraction, multiplication, and square"],
        ["Maximum order", "Quadratic"],
        ["Interactions", "Pairwise only"],
        ["Forbidden forms", "Ratios except as explicitly disallowed, log, exp, trigonometric functions, cubic and higher powers"],
        ["LLM output", "One expression structure; no numerical coefficients; no explanation"],
        ["Coefficient fitting", "OLS on training data"],
        ["Primary selection", "Training five-fold CV plus parsimony and diagnostics"],
        ["Test usage", "Post hoc after the structure freeze"],
    ],
    [2200, 7160],
    font_size=8.1,
)
for idx, stage_name in enumerate(["M0", "M1", "M2", "M3", "M4"], start=1):
    prompt = (ROOT / "llm_repeated" / "prompts" / f"{stage_name}.txt").read_text(encoding="utf-8")
    add_prompt_box(idx, stage_name, prompt)
add_callout(
    "Physical-knowledge implementation",
    "Knowledge was encoded through closure-aware variables, restricted operator sets, staged expansion from main effects to curvature and selected interactions, and post hoc coefficient diagnostics. No coefficient sign was forced.",
)

add_heading("Supplementary Note S5. Stage-wise candidate generation and M0-M4 evolution", 1)
add_body(
    "The stage ladder was designed to separate increasingly complex physical hypotheses. M0 tests linear substitution effects. M1 adds single-variable curvature. M2 opens selected metal-halide and mixed-halide interactions. M3 introduces restricted A-site terms. Final M4 is a new diagnostic freeze obtained after removing unstable and collinearity-heavy terms from the larger late-stage candidates."
)
add_table(
    "Table S8. Stage-wise admissible term space.",
    ["Stage", "Admissible terms", "Purpose"],
    [
        ["M0", "Sn, Br, Cl", "Linear baseline"],
        ["M1", "M0 + Sn², Br², Cl²", "Single-variable curvature"],
        ["M2", "M1 + SnBr, SnCl, BrCl", "B-X and mixed-halide coupling"],
        ["M3", "M2 + Cs, Cs², CsSn, CsBr, CsCl", "Restricted A-site testing"],
        ["Final M4", "Sn, Br, Cl, Cs, Sn², Br², CsSn, CsCl", "Frozen compact formula after diagnostics"],
    ],
    [900, 5200, 3260],
    font_size=7.9,
)
add_table(
    "Table S9. Stage-wise performance. Test values are post hoc and were not used to select the final structure.",
    ["Model", "Terms", "CV RMSE", "Test RMSE", "Role"],
    [
        ["M0", "3", "0.087733", "0.085887", "Linear baseline"],
        ["M1", "6", "0.063939", "0.071660", "Main effects + squares"],
        ["M2", "9", "0.066548", "0.066416", "Full anion quadratic"],
        ["M3-full", "14", "0.058333", "0.054620", "Expressive Cs interaction model"],
        ["Historical 13-term M4", "13", "0.059460", "0.053835", "Pre-final parsimony candidate"],
        ["Final frozen 8-term M4", "8", "0.057532", "0.060613", "Only final model in the revision"],
    ],
    [1900, 850, 1350, 1350, 3910],
    font_size=8.0,
)
add_figure(
    "figure_s4_stage_performance.png",
    "Figure S4. Cross-validation RMSE versus model complexity. The post hoc test curve is shown for transparency but was not used for the freeze decision. Final M4 attains the lowest reported fixed-CV error among the named stage models while using eight terms.",
)
add_body(
    "The historical M3 and 13-term M4 workbooks predate the diagnostic freeze and are retained as an audit trail. Their lower post hoc test errors do not override the training-CV and coefficient-stability rules. The final formula was refitted after the eight-term structure was locked."
)

add_heading("Supplementary Note S6. Final M4 freezing rule and model nomenclature", 1)
add_body(
    "The freeze rule combined: (i) training five-fold CV qualification; (ii) constrained parsimony; (iii) VIF and confidence-interval checks; (iv) 1000-resample bootstrap sign stability; (v) Cl² nested-model diagnostics; and (vi) subgroup and group-aware sensitivity. Test RMSE was not an optimization target."
)
add_table(
    "Table S10. Model nomenclature and final-version locking.",
    ["Locked name", "Term set", "Role", "Final?"],
    [
        ["Final frozen 8-term M4", "Sn, Br, Cl, Cs, Sn², Br², CsSn, CsCl", "Revision endpoint", "Yes"],
        ["Historical 13-term M4", "Includes Cl², Cs², SnBr, SnCl, CsBr, CsCl", "Pre-final candidate", "No"],
        ["M3-full", "14-term Cs interaction polynomial", "Expressive late-stage reference", "No"],
        ["LLM-stable compact candidate", "Sn, Br, Cl, Cs, Sn², Br², CsSn, CsBr", "30-run stability reference", "No"],
        ["Anion-only exhaustive best", "Sn, Br, Cl, Sn², Br², SnBr", "Deterministic ablation", "No"],
        ["Anion-only main + quadratic", "Sn, Br, Cl, Sn², Br², Cl²", "Historical diagnostic candidate", "No"],
    ],
    [2300, 3450, 2450, 1160],
    font_size=7.7,
)
add_callout(
    "Why the historical model is not final",
    "The 13-term candidate has a lower test RMSE (0.053835 eV) than the frozen 8-term model, but it was not selected on the test set. It contains diagnostically weak and collinearity-heavy terms. Calling it final would contradict the declared training-only selection protocol.",
    fill=CAUTION_FILL,
)

page_break()
add_heading("Supplementary Note S7. Coefficient diagnostics, VIF, confidence intervals, and bootstrap", 1)
add_body(
    "Final M4 was fitted by OLS on all 518 training rows after the structure freeze. OLS intervals use the standard linear-model covariance estimate. The bootstrap used 1000 row-resampling fits with seed 2026. Sign stability is the fraction of bootstrap coefficients with the same sign as the full-training estimate."
)
ols_rows = []
boot_rows = []
for _, row in coeff.iterrows():
    pval = float(row["p-value"])
    ptext = "<1e-300" if pval == 0 else f"{pval:.2e}"
    vif = "-" if pd.isna(row["VIF"]) else f'{float(row["VIF"]):.2f}'
    ols_rows.append(
        [
            row["Term"],
            f'{float(row["Coefficient"]):.6f}',
            f'{float(row["SE"]):.6f}',
            ptext,
            f'[{float(row["OLS CI low"]):.6f}, {float(row["OLS CI high"]):.6f}]',
            vif,
        ]
    )
    boot_rows.append(
        [
            row["Term"],
            f'{float(row["Bootstrap mean"]):.6f}',
            f'{float(row["Bootstrap SD"]):.6f}',
            f'[{float(row["Bootstrap CI low"]):.6f}, {float(row["Bootstrap CI high"]):.6f}]',
            f'{float(row["Bootstrap sign stability"]):.3f}',
        ]
    )
add_table(
    "Table S11A. Final M4 OLS coefficient statistics and VIF.",
    ["Term", "Coefficient", "SE", "p value", "OLS 95% CI", "VIF"],
    ols_rows,
    [900, 1100, 1000, 1650, 3200, 1510],
    font_size=7.6,
)
add_table(
    "Table S11B. Final M4 bootstrap coefficient statistics.",
    ["Term", "Bootstrap mean", "Bootstrap SD", "Bootstrap 95% CI", "Sign stability"],
    boot_rows,
    [900, 1500, 1300, 3400, 2260],
    font_size=7.6,
)
add_body(
    "All retained slope intervals exclude zero. Sign stability is 1.000 for seven slopes and 0.999 for CsCl. Sn and Sn² have VIF values near 15, and Br and Br² near 9, reflecting expected dependence between raw polynomial terms. The interaction terms remain low-VIF. These values are materially lower than the maximum VIF above 3000 in the 14-term full polynomial."
)
cl2_map = {row["Metric"]: row for _, row in cl2.iterrows()}
add_table(
    "Table S12. Removed-term diagnostic for adding Cl² to Final M4.",
    ["Diagnostic", "Final M4", "Final M4 + Cl²", "Interpretation"],
    [
        ["CV RMSE", "0.057532", "0.058256", "Worse with Cl²"],
        ["Test RMSE", "0.060613", "0.062508", "Post hoc error also worse"],
        ["AIC", "-3003.783", "-3003.233", "No support for added term"],
        ["BIC", "-2965.533", "-2960.733", "Penalizes added term"],
        ["Maximum VIF", "15.786", "16.177", "Slightly higher"],
        ["Cl² coefficient", "-", "0.103629", "Small relative to uncertainty"],
        ["Cl² OLS 95% CI", "-", "[-0.066945, 0.274202]", "Crosses zero"],
        ["Cl² bootstrap 95% CI", "-", "[-0.411010, 0.505612]", "Crosses zero"],
        ["Cl² sign stability", "-", "0.6034", "Insufficient"],
        ["Nested F-test p value", "-", "0.2332", "Added term not supported"],
    ],
    [2300, 1700, 2100, 3260],
    font_size=7.8,
)
add_figure(
    "figure_s5_final_diagnostics.png",
    "Figure S5. Final M4 coefficient diagnostics. (a) OLS coefficients and 95% confidence intervals. (b) VIF for raw polynomial terms. (c) Bootstrap sign stability across 1000 resamples.",
)
add_body(
    "A 100-repeat five-fold CV audit gave mean RMSE 0.058056 eV with between-partition SD 0.008714 eV, close to the fixed-fold value 0.057532 eV. Final training RMSE, MAE, and R² were 0.054108 eV, 0.039261 eV, and 0.969382. Held-out test RMSE, MAE, median absolute error, maximum absolute error, and R² were 0.060613 eV, 0.046760 eV, 0.036645 eV, 0.183551 eV, and 0.976545."
)

add_heading("Supplementary Note S8. MA and A-site encoding sensitivity", 1)
add_body(
    "MA was not removed because it is physically irrelevant. Under FA + MA + Cs = 1, all three A-site fractions cannot enter an intercept-containing linear model without exact dependence. FA was treated as the omitted baseline. The sensitivity analysis compared the nine-term Sn-Br-Cl core, isolated Cs or MA main effects, joint Cs and MA main effects, and the full Cs interaction structure."
)
ma_test = {
    "Sn-Br-Cl core": 0.066416,
    "Sn-Br-Cl + Cs": 0.065252,
    "Sn-Br-Cl + Cs interactions": 0.054620,
    "Sn-Br-Cl + MA": 0.065933,
    "Sn-Br-Cl + Cs + MA": 0.063302,
}
ma_rows = []
for _, row in ma.sort_values("CV rank").iterrows():
    ma_rows.append(
        [
            row["Model"],
            row["Encoding"],
            int(row["k"]),
            f'{float(row["CV RMSE"]):.6f}',
            f'{ma_test[row["Model"]]:.6f}',
            "Training CV primary; test post hoc",
        ]
    )
ma_rows.append(
    ["Final frozen M4", "Cs main + CsSn + CsCl", 8, "0.057532", "0.060613", "Diagnostic freeze"]
)
add_table(
    "Table S13. A-site encoding sensitivity.",
    ["Model", "A-site encoding", "k", "CV RMSE", "Test RMSE", "Interpretation"],
    ma_rows,
    [2300, 2400, 500, 1150, 1150, 1860],
    font_size=7.4,
)
add_body(
    "The core CV RMSE is 0.066548 eV. Adding Cs alone leaves the error unchanged at 0.066566 eV. MA alone gives 0.065590 eV, and Cs + MA main effects give 0.064405 eV. The Cs interaction model reaches 0.058333 eV. Thus the available A-site signal is mainly conditional, not a large isolated linear MA or Cs effect. The final compact model retains Cs, CsSn, and CsCl. MA-rich systems remain part of the applicability-domain caution because explicit MA did not show a systematic compact-model advantage."
)

page_break()
add_heading("Supplementary Note S9. Non-LLM symbolic regression and statistical controls", 1)
add_body(
    "Conventional controls were evaluated on the same immutable train/test files. Their purpose is not to prove universal superiority of one search algorithm. They test whether compact low-order structures are reachable without an LLM and whether the final freeze can be justified by deterministic or stochastic alternatives."
)
add_table(
    "Table S14. Compact GPLearn baseline.",
    ["Setting or result", "Archived value"],
    [
        ["Variables", "Sn, Br, Cl, Cs, MA"],
        ["Function set", "add, subtract, multiply"],
        ["Population / generations", "300 / 12"],
        ["Tournament size", "20"],
        ["Parsimony grid", "1e-6, 1e-5, 1e-4"],
        ["Selection", "Minimum mean five-fold training CV RMSE"],
        ["CV RMSE ± SD", f"{gplearn['cv_rmse']:.6f} ± {gplearn['cv_sd']:.6f} eV"],
        ["Independent full-training refits", f"{gplearn['n_refit_seeds']} ({gplearn['unique_programs']} unique programs)"],
        ["Test RMSE mean ± SD", f"{gplearn['test_rmse_mean']:.6f} ± {gplearn['test_rmse_sd']:.6f} eV"],
        ["Test RMSE range", f"{gplearn['test_rmse_min']:.6f}-{gplearn['test_rmse_max']:.6f} eV"],
        ["Best post hoc program", gplearn["best_posthoc_program"]],
    ],
    [3000, 6360],
    font_size=7.9,
)
add_callout(
    "GPLearn scope",
    "The operator set and search budget were deliberately compact. This result is a limited-budget control, not a universal statement about genetic programming.",
)
pysr_rows = []
for key, label in [("P_polynomial", "Polynomial operators"), ("R_rich", "Richer operators")]:
    item = pysr[key]
    complexities = [run["best_complexity"] for run in item["per_seed"]]
    pysr_rows.append(
        [
            label,
            item["n_seeds"],
            f'{item["best_test_rmse_mean"]:.4f} ± {item["best_test_rmse_std"]:.4f}',
            f'{item["best_test_rmse_min"]:.4f}',
            item["n_unique_best_structures"],
            f"{min(complexities)}-{max(complexities)}",
        ]
    )
add_table(
    "Table S15. PySR repeated-run results.",
    ["Search space", "Seeds", "Mean best test RMSE", "Minimum", "Unique best structures", "Complexity range"],
    pysr_rows,
    [2100, 700, 2100, 1100, 1800, 1560],
    font_size=7.8,
)
add_figure(
    "figure_s7_symbolic_baselines.png",
    "Figure S7. Accuracy-complexity reference for symbolic approaches. Complexity definitions are method-specific and are used only as an approximate visual reference.",
)
add_table(
    "Table S16. Exhaustive search over all 511 non-empty subsets of the nine-term anion polynomial dictionary.",
    ["Best k", "Best terms", "CV RMSE", "Test RMSE", "Pareto status"],
    [
        ["1", "Cl", "0.225271", "0.241747", "Yes"],
        ["2", "Br, Cl", "0.134708", "0.144643", "Yes"],
        ["3", "Sn, Br, Cl", "0.087733", "0.085887", "Yes"],
        ["4", "Sn, Br, Cl, Sn²", "0.068229", "0.068946", "Yes"],
        ["5", "Sn, Br, Cl, Sn², Br²", "0.063099", "0.071471", "Yes"],
        ["6", "Sn, Br, Cl, Sn², Br², SnBr", "0.062853", "0.066711", "Global best"],
        ["7", "Sn, Br, Cl, Sn², Br², SnBr, BrCl", "0.063671", "0.067048", "No"],
        ["8", "Sn, Br, Cl, Sn², Br², Cl², SnBr, BrCl", "0.063611", "0.068587", "No"],
        ["9", "All nine anion terms", "0.066548", "0.066416", "No"],
    ],
    [750, 4200, 1350, 1350, 1710],
    font_size=7.6,
)
add_body(
    "The deterministic optimum is a six-term anion model and excludes Cl². It is a strong compact baseline but does not include the A-site interactions that lower the final M4 CV error. Exhaustive search therefore shows that low-order polynomial structure is not unique to the LLM workflow while also weakening the case for retaining Cl²."
)
reg_rows = []
for _, row in regularization.iterrows():
    reg_rows.append(
        [
            row["Method"],
            row["Rule"],
            int(row["k"]),
            row["Selected terms"],
            f'{float(row["CV RMSE"]):.6f}',
            "Yes" if row["Cl2 zero"] == "Yes" else "No",
        ]
    )
reg_rows.extend(
    [
        ["Stepwise", "AIC", 7, "Cl, Br, Sn, Sn², Br², SnCl, SnBr", "n/a", "Yes"],
        ["Stepwise", "BIC", 7, "Cl, Br, Sn, Sn², Br², SnCl, SnBr", "n/a", "Yes"],
        ["All subsets", "AIC", 9, "All nine anion terms", "n/a", "No"],
        ["All subsets", "BIC", 9, "All nine anion terms", "n/a", "No"],
    ]
)
add_table(
    "Table S17. Regularization and information-criterion controls.",
    ["Method", "Rule", "k", "Selected terms", "CV RMSE", "Cl² removed?"],
    reg_rows,
    [1350, 1500, 500, 3900, 1100, 1010],
    font_size=7.2,
)
add_body(
    "Minimum-CV LASSO and Elastic Net keep all nine terms because their selected penalties are effectively zero. Their 1-SE rules remove Cl² and BrCl but retain SnBr and SnCl. Stepwise AIC and BIC select the same seven-term interaction model, whereas exhaustive AIC/BIC favor the full nine-term model. These controls support pruning as a reasonable objective but do not reproduce the exact final support."
)
add_callout(
    "Control conclusion",
    "Conventional methods can reach low-order polynomial structures. The LLM-SR workflow is best interpreted as a controlled hypothesis-generation route that combines domain constraints, parsimony, and diagnostics, not as a universally superior optimizer.",
)

add_heading("Supplementary Note S10. Repeated LLM runs and cross-model stochasticity", 1)
add_body(
    "Stochasticity was examined in three ways: 30 direct M4 prompt calls, 30 complete M0-M4 workflows, and repeated Qwen and DeepSeek workflows at three temperatures. Repeated runs were used to assess the stability of the M4-family backbone, not to choose the final formula by majority vote."
)
add_table(
    "Table S18. Repeated LLM run settings.",
    ["Experiment", "Model/provider identifier", "Runs", "Generation controls", "Local evaluation"],
    [
        ["Direct M4 prompts", "Archived Codex/gpt-5.5 identifier", "30", "Temperature and seed not exposed; isolated sessions", "Term normalization only"],
        ["Complete M0-M4 workflows", "Archived Codex/gpt-5.5 identifier", "30", "16 candidates per stage; reasoning effort low", "CV seed 20260607; 2400 accepted candidates"],
        ["Qwen workflows", "qwen-plus", "10 per T", "T = 0.2, 0.7, 1.0; top_p = 1; max_tokens = 3000", "CV seed 20260607"],
        ["DeepSeek workflows", "deepseek-v4-flash", "10 per T", "T = 0.2, 0.7, 1.0; top_p = 1; max_tokens = 3000", "CV seed 20260607"],
    ],
    [1800, 2100, 900, 2850, 1710],
    font_size=7.3,
)
direct_freq = {
    "Sn": 30, "Br": 30, "Cl": 30, "Cs": 30, "Sn²": 30, "Br²": 28, "Cl²": 24,
    "Cs²": 12, "SnBr": 30, "SnCl": 30, "BrCl": 2, "CsSn": 3, "CsBr": 29, "CsCl": 9,
}
full_freq = {
    "Sn": 30, "Br": 30, "Cl": 30, "Cs": 30, "Sn²": 30, "Br²": 30, "Cl²": 0,
    "Cs²": 0, "SnBr": 0, "SnCl": 0, "BrCl": 0, "CsSn": 30, "CsBr": 30, "CsCl": 0,
}
term_rows = [
    [term, direct_freq[term], f"{direct_freq[term] / 30 * 100:.1f}%", full_freq[term], f"{full_freq[term] / 30 * 100:.1f}%"]
    for term in direct_freq
]
add_table(
    "Table S19. Term frequency in direct M4 outputs and selected models from 30 complete workflows.",
    ["Term", "Direct count", "Direct %", "Full-workflow count", "Full-workflow %"],
    term_rows,
    [1200, 1700, 1400, 2500, 2560],
    font_size=7.7,
)
add_body(
    f"The 30 direct M4 outputs contained {direct_summary['unique_canonical_structures']} normalized structures. The complete workflows converged to one eight-term CsBr-containing candidate in 30/30 runs, with CV RMSE {repeated_summary['cv_rmse']['mean']:.6f} eV and test RMSE {repeated_summary['test_rmse']['mean']:.6f} eV. This establishes empirical stability under the archived configuration, not deterministic seed-controlled reproducibility."
)
cross_rows = [
    [
        row["provider"],
        f'{float(row["T"]):.1f}',
        int(row["n"]),
        int(row["unique_structures"]),
        f'{float(row["cv_median"]):.6f}',
        f'{float(row["test_median"]):.6f}',
        f'{float(row["test_iqr"]):.6f}',
        f'{float(row["test_min"]):.6f}-{float(row["test_max"]):.6f}',
    ]
    for _, row in external_llm.iterrows()
]
add_table(
    "Table S20. Cross-provider repeated workflow summary.",
    ["Provider", "T", "n", "Unique structures", "Median CV", "Median test", "Test IQR", "Test range"],
    cross_rows,
    [1100, 500, 500, 1500, 1300, 1300, 1200, 1960],
    font_size=7.2,
)
add_figure(
    "figure_s6_llm_stochasticity.png",
    "Figure S6. LLM stochasticity. (a) Term occurrence in 30 independent direct M4 prompt outputs. Dark cells indicate inclusion. (b) Median held-out error for repeated Qwen and DeepSeek workflows; error bars show half the interquartile range.",
)
add_body(
    "CsBr is frequent in repeated generation and is the terminal term in all 30 complete-workflow selected candidates. CsCl is less frequent in direct generation. The final revision nevertheless retains CsCl because the freeze decision was made after local OLS fitting, CV, confidence intervals, bootstrap sign stability, and interaction-specific diagnostics. Majority vote over generated structures was not the selection rule."
)

page_break()
add_heading("Supplementary Note S11. Black-box baselines and SHAP cross-check", 1)
add_body(
    "Gaussian-process regression and tree ensembles provide predictive reference points using the same compact composition inputs. They are not equation-discovery methods, but they test the accuracy cost of restricting the response to eight explicit global terms."
)
add_table(
    "Table S21. Black-box and final analytical model performance.",
    ["Method", "Inputs/model", "CV RMSE", "Test RMSE", "Test MAE", "Test R²"],
    [
        ["Gaussian process", "Sn, Br, Cl, Cs, MA; Matérn-1.5", "0.051320", "0.048604", "0.037050", "0.984919"],
        ["GBRT", "Sn, Br, Cl, Cs, MA", "0.053247", "0.052021", "0.040229", "0.982724"],
        ["Random forest", "Sn, Br, Cl, Cs, MA", "0.059683", "0.055863", "n/a", "0.980078"],
        ["XGBoost", "Sn, Br, Cl, Cs, MA", "0.051267", "0.060565", "n/a", "0.976583"],
        ["Final M4", "Eight-term OLS equation", "0.057532", "0.060613", "0.046760", "0.976545"],
    ],
    [1700, 2900, 1250, 1250, 1100, 1160],
    font_size=7.8,
)
add_figure(
    "figure_s8_gbrt_shap_beeswarm.png",
    "Figure S8. GBRT SHAP beeswarm. Br, Sn, and Cl dominate the response; Cs and MA are smaller corrections.",
    width=5.7,
)
add_figure(
    "figure_s9_rf_xgb_shap_beeswarm.png",
    "Figure S9. SHAP beeswarm plots for random forest and XGBoost. Both reproduce the same dominant feature hierarchy and response directions as GBRT.",
)
add_figure(
    "figure_s10_shap_interactions.png",
    "Figure S10. GBRT SHAP interaction matrix. The diagonal represents main effects; the largest off-diagonal interaction is Br-Cl.",
    width=5.5,
)
add_table(
    "Table S22. Symbolic coefficient trends versus SHAP trends.",
    ["Feature or pair", "Final M4 representation", "SHAP evidence", "Cross-check"],
    [
        ["Sn", "Negative main; positive Sn² curvature", "GBRT mean |SHAP| 0.0966; dependence decreases with Sn", "Consistent"],
        ["Br", "Positive main and Br² curvature", "GBRT 0.1078; RF 0.1087; XGB 0.0984", "Consistent dominant effect"],
        ["Cl", "Strong positive main", "GBRT 0.0813; sparse high positive responses", "Consistent but data-sparse"],
        ["Cs", "Small positive main; negative CsSn and CsCl", "GBRT 0.0128; RF 0.0126; XGB 0.0127", "Small conditional effect"],
        ["MA", "Not explicit in final formula", "GBRT 0.0121; RF 0.0099; XGB 0.0120", "Small omitted-baseline correction"],
        ["Br-Cl", "No global term retained", "Top interaction: 0.00725 GBRT; 0.00841 RF; 0.01651 XGB", "Localized nonlinearity omitted by compact formula"],
        ["Sn-Cs", "CsSn retained", "GBRT interaction 0.00436", "Consistent coupling"],
    ],
    [1300, 2400, 3300, 2360],
    font_size=7.3,
)
add_body(
    "SHAP supports the dominant roles of halide and Sn substitution while identifying Br-Cl as a localized nonlinear region better captured by black-box models. The absence of BrCl from Final M4 is therefore a stability-parsimony decision, not a claim that Br-Cl coupling is physically absent."
)

add_heading("Supplementary Note S12. Group-aware splitting and applicability domain", 1)
add_body(
    "Group-aware analysis used the 507-row strict-closure subset and repeated the full candidate-selection workflow. Composition-family group shuffle tests generalization across combined A/B/X families. Halide and A-site leave-one-group-out tests deliberately remove entire chemistry regimes from training. These are sensitivity analyses and did not alter the frozen random-split model."
)
group_rows = [
    [
        row["group_strategy"],
        int(row["n_splits"]),
        f'{row["mean_test_rmse"]:.6f}',
        f'{row["sd_test_rmse"]:.6f}',
        f'{row["median_test_rmse"]:.6f}',
        f'{row["worst_test_rmse"]:.6f}',
        f'{row["mean_jaccard_to_M4"]:.3f}',
    ]
    for _, row in group_summary.iterrows()
]
add_table(
    "Table S23. Group-aware split performance summary.",
    ["Strategy", "Splits", "Mean RMSE", "SD", "Median", "Worst", "Mean Jaccard"],
    group_rows,
    [2600, 700, 1250, 1050, 1200, 1200, 1360],
    font_size=7.5,
)
detailed = group_heldout[group_heldout.group_strategy.isin(["halide_logo", "a_site_logo"])]
heldout_rows = [
    [
        row["group_strategy"],
        str(row["heldout_groups"]).replace("_", " "),
        f'{row["mean_test_rmse"]:.6f}',
        f'{row["mean_test_mae"]:.6f}',
        f'{row["mean_jaccard_to_M4"]:.3f}',
    ]
    for _, row in detailed.iterrows()
]
add_table(
    "Table S24A. Leave-one-group-out results.",
    ["Strategy", "Held-out group", "RMSE", "MAE", "Jaccard to M4"],
    heldout_rows,
    [1850, 2300, 1500, 1500, 2210],
    font_size=7.6,
)
selected_terms = ["Sn", "Br", "Cl", "Sn2", "Br2", "Cl2", "Cs", "CsSn", "CsBr", "CsCl", "BrCl"]
freq_rows = []
for strategy in ["composition_family_group_shuffle", "halide_logo", "a_site_logo"]:
    for term in selected_terms:
        row = group_terms[(group_terms.group_strategy == strategy) & (group_terms.term == term)]
        if row.empty:
            continue
        item = row.iloc[0]
        freq_rows.append(
            [strategy, term, int(item["count_selected"]), f'{float(item["frequency_percentage"]):.1f}%']
        )
add_table(
    "Table S24B. Selected-term frequencies under group-aware workflows.",
    ["Strategy", "Term", "Selected count", "Frequency"],
    freq_rows,
    [3600, 1500, 1900, 2360],
    font_size=7.4,
)
add_figure(
    "figure_s11_group_term_frequency.png",
    "Figure S11. Term-selection frequency under composition-family, halide-LOGO, and A-site-LOGO workflows.",
)
add_figure(
    "figure_s12_group_heldout_error.png",
    "Figure S12. Held-out-group errors. Removing the full Cl-containing regime produces the dominant failure case; Cs-rich holdout is the strongest A-site boundary.",
)
add_body(
    "Composition-family splits give mean RMSE 0.08844 eV and median 0.07496 eV. Holding out mixed-halide, I-rich, or Br-rich groups gives 0.05564, 0.06620, and 0.07259 eV, whereas holding out all Cl-containing compositions gives 0.88940 eV. A-site LOGO gives 0.02612 eV for mixed-A, 0.05691 eV for FA-rich, 0.08501 eV for MA-rich, and 0.17382 eV for Cs-rich."
)
add_callout(
    "Applicability-domain conclusion",
    "Final M4 is an interpolation surrogate for the dominant represented manifold. Strict halide extrapolation, especially into Cl-containing families without Cl examples in training, and isolated Cs-rich families require additional data or a nonlinear model.",
    fill=CAUTION_FILL,
)

page_break()
add_heading("Supplementary Note S13. Independent external-validation status", 1)
add_body(
    "The 92-row held-out workbook is drawn from the same frozen database snapshot as the training data. It is a legitimate non-selection test set but not an independent literature validation set. The project archive does not contain a verified external panel with composition provenance, measurement conditions, deduplication against the source database, and harmonized optical-gap definitions."
)
add_table(
    "Table S25. External-validation evidence status and required completion protocol.",
    ["Evidence item", "Current archive status", "Submission-ready requirement"],
    [
        ["Independent literature compositions", "Not available", "Assemble a composition-resolved external set"],
        ["Proof of non-overlap with source database", "Not available", "Deduplicate by composition, citation, and measurement record"],
        ["Temperature, phase, and measurement method", "Not harmonized", "Archive metadata and uncertainty"],
        ["Final M4 predictions", "Not generated for an independent set", "Apply frozen formula without refitting"],
        ["External RMSE, MAE, median AE, max AE", "Not reported", "Report overall and subgroup metrics"],
        ["Parity and absolute-error figures", "Not produced", "Generate only after the dataset is verified"],
        ["Model-selection role", "Must remain none", "Use external data for transferability reporting only"],
    ],
    [2800, 2800, 3760],
    font_size=7.8,
)
add_callout(
    "No fabricated validation",
    "This SI intentionally omits an external parity plot and sample list because the required source-grounded dataset is absent. The strongest current evidence is the immutable held-out split plus the group-aware sensitivity analysis.",
    fill=RED_FILL,
    color=RED,
)

add_heading("Supplementary Note S14. Composition-guided design demonstration", 1)
add_body(
    "The closed-form model can be evaluated over dense composition grids without retraining. Figure S13 shows two response slices. Colored regions are restricted by distance to the training manifold; white regions are withheld rather than presented as valid predictions."
)
add_figure(
    "figure_s13_design_maps.png",
    "Figure S13. Band-gap landscapes generated by Final M4. Colored points are within Euclidean distance 0.25 of a training composition in the Sn-Br-Cl-Cs space. Hatching marks the illustrative 1.60-1.80 eV target window.",
)
add_body(
    f"An additional screen generated {summary['design_screen']['generated']:,} local perturbations around archived training compositions with seed {summary['design_screen']['seed']}. A-site and X-site fractions were perturbed on their simplices, Sn was locally perturbed, and Pb was set by closure. Candidates were required to have predicted Eg between 1.60 and 1.80 eV, nearest-training distance no greater than 0.15 in Sn-Br-Cl-Cs-MA space, Cs no greater than 0.50, and Cl no greater than 0.15. This left {summary['design_screen']['eligible']:,} illustrative points. Twelve compositionally separated examples are reported below."
)
design_rows = []
for _, row in design.iterrows():
    design_rows.append(
        [
            row["Candidate"],
            f'{row["FA"]:.3f}/{row["MA"]:.3f}/{row["Cs"]:.3f}',
            f'{row["Pb"]:.3f}/{row["Sn"]:.3f}',
            f'{row["I"]:.3f}/{row["Br"]:.3f}/{row["Cl"]:.3f}',
            f'{row["Eg_pred"]:.3f}',
            f'{row["nearest_train_distance"]:.3f}',
        ]
    )
add_table(
    "Table S26. Illustrative locally constrained candidates in the 1.60-1.80 eV window. Fraction strings are FA/MA/Cs, Pb/Sn, and I/Br/Cl.",
    ["Candidate", "A site", "B site", "X site", "Pred. Eg", "Nearest d"],
    design_rows,
    [1000, 2000, 1500, 2300, 1200, 1360],
    font_size=7.7,
)
add_body(
    "These rows are model-guided examples, not discovered materials. They do not include phase stability, defect tolerance, synthesis feasibility, processing conditions, or uncertainty in the experimental target. Candidate ranking should therefore be followed by nonlinear local refinement and experimental or physics-based validation."
)

page_break()
add_heading("Supplementary Note S15. Reproducibility package and file manifest", 1)
add_body(
    "The repository contains the fixed data files, prompts, repeated-run records, statistical diagnostics, black-box outputs, group-aware manifests, and scripts listed below. The release is organized for public audit; a versioned GitHub URL and Zenodo DOI should be added to the manuscript metadata after publication."
)
manifest_rows = [
    ["data/dataset_610_snapshot.xlsx", "610-row record-level snapshot plus auxiliary sheets", "Dataset audit", "Present"],
    ["data/train_518.xlsx", "Immutable 518-row training workbook", "All model fitting and selection", "Present"],
    ["data/test_92.xlsx", "Immutable 92-row held-out workbook", "Post hoc evaluation", "Present"],
    ["archive/historical_model_stage_history.xlsx", "Historical stage formulas", "Stage audit", "Present"],
    ["llm_repeated/prompts/M0.txt ... M4.txt", "Exact direct prompt templates", "Prompt reproducibility", "Present"],
    ["final_m4_diagnostics/final_M4_coefficients_CI_bootstrap.csv", "OLS, CI, bootstrap, and VIF statistics", "Final diagnostics", "Present"],
    ["final_m4_diagnostics/Cl2_retention_diagnostic.csv", "Nested Cl² comparison", "Removed-term evidence", "Present"],
    ["ma_asite_sensitivity/ma_asite_encoding_sensitivity.xlsx", "A-site encoding comparison", "MA sensitivity", "Present"],
    ["baselines/exhaustive_polynomial_search/exhaustive_polynomial_search.xlsx", "All 511 anion models and Pareto tables", "Deterministic control", "Present"],
    ["baselines/regularization_aic_bic/*.xlsx,csv", "Regularization and information criteria", "Selection controls", "Present"],
    ["baselines/pysr/", "Five-seed PySR outputs and Pareto files", "Symbolic baseline", "Present"],
    ["scripts/reproduce_gplearn.py", "Compact GPLearn experiment", "Symbolic baseline", "Present"],
    ["baselines/gplearn/", "GPLearn CV and seed outputs", "Symbolic baseline", "Present"],
    ["llm_repeated/raw_outputs/", "Thirty direct M4 raw/normalized outputs", "Stochasticity", "Present"],
    ["llm_repeated/repeated_runs_30/", "Thirty complete M0-M4 workflows", "Stochasticity and audit", "Present"],
    ["llm_repeated/external_api_runs/", "Qwen/DeepSeek runs, manifests, and summaries", "Cross-model stability", "Present"],
    ["baselines/gp_learning/", "Gaussian-process selection and predictions", "Predictive benchmark", "Present"],
    ["blackbox_shap/gbrt/", "GBRT model selection, SHAP arrays, and figures", "Interpretability", "Present"],
    ["blackbox_shap/rf_xgboost/", "RF/XGBoost metrics and SHAP outputs", "Interpretability", "Present"],
    ["group_aware_sensitivity/scripts,src,outputs", "Validated group splits, workflows, tests, and figures", "Applicability domain", "Present"],
    ["final_m4/", "Final freeze summary and formula records", "Version lock", "Present"],
    ["scripts/publication/prepare_si_assets.py", "SI audit tables, design screen, and figures", "This SI", "Present"],
    ["scripts/publication/build_si_docx.py", "Deterministic DOCX builder", "This SI", "Present"],
    ["environment.yml and requirements.txt", "Complete software environment", "Portable rerun", "Present"],
    ["Public GitHub URL", "Versioned public repository", "Access", "Not present"],
    ["Zenodo DOI", "Immutable archival release", "Citation", "Not present"],
]
add_table(
    "Table S27. Reproducibility package manifest.",
    ["File or folder", "Description", "Used for", "Status"],
    manifest_rows,
    [3300, 3000, 1900, 1160],
    font_size=7.0,
)
add_body(
    "The three principal SHA-256 checksums are: `data/dataset_610_snapshot.xlsx`, "
    f"{summary['checksums']['data/dataset_610_snapshot.xlsx']}; `data/train_518.xlsx`, "
    f"{summary['checksums']['data/train_518.xlsx']}; `data/test_92.xlsx`, "
    f"{summary['checksums']['data/test_92.xlsx']}. The repeated-workflow manifests also record prompt hashes, model identifiers, CV seeds, candidate counts, Python, NumPy, and pandas versions."
)
add_heading("Reproduction sequence", 2)
for label, text in [
    ("1", "Verify the three data-file checksums and use the immutable training/test workbooks."),
    ("2", "Read the exact M0-M4 templates and candidate schemas in `llm_repeated`."),
    ("3", "Run local OLS and five-fold CV; do not use the held-out workbook for candidate selection."),
    ("4", "Regenerate final coefficient diagnostics from the VIF/CI/bootstrap workflow."),
    ("5", "Regenerate non-LLM, black-box, SHAP, and group-aware controls from their archived scripts and manifests."),
    ("6", "Use `scripts/publication/prepare_si_assets.py` and `scripts/publication/build_si_docx.py` to reproduce this document."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.20)
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(f"Step {label}: "), size=9.5, bold=True, color=NAVY)
    set_run_font(p.add_run(text), size=9.5)

add_heading("Data Availability Statement", 2)
add_body(
    "The frozen local dataset snapshot, immutable train/test workbooks, exact prompts, repeated-run logs, benchmark outputs, SHAP summaries, group-aware split manifests, and analysis scripts are contained in the accompanying project archive. A public GitHub release and immutable Zenodo archive should be created before submission. No unverified public URL or DOI is asserted here."
)

add_heading("Supplementary References", 1)
references = [
    "1. Jacobsson, T. J.; Hultqvist, A.; García-Fernández, A.; et al. An Open-Access Database and Analysis Tool for Perovskite Solar Cells Based on the FAIR Data Principles. Nat. Energy 2022, 7, 107-115. https://doi.org/10.1038/s41560-021-00941-3.",
    "2. Schmidt, M.; Lipson, H. Distilling Free-Form Natural Laws from Experimental Data. Science 2009, 324, 81-85. https://doi.org/10.1126/science.1165893.",
    "3. Cranmer, M. Interpretable Machine Learning for Science with PySR and SymbolicRegression.jl. arXiv 2023, arXiv:2305.01582. https://doi.org/10.48550/arXiv.2305.01582.",
    "4. Shojaee, P.; Meidani, K.; Gupta, S.; Farimani, A. B.; Reddy, C. K. LLM-SR: Scientific Equation Discovery via Programming with Large Language Models. arXiv 2024, arXiv:2404.18400. https://doi.org/10.48550/arXiv.2404.18400.",
    "5. Breiman, L. Random Forests. Mach. Learn. 2001, 45, 5-32. https://doi.org/10.1023/A:1010933404324.",
    "6. Friedman, J. H. Greedy Function Approximation: A Gradient Boosting Machine. Ann. Stat. 2001, 29, 1189-1232. https://doi.org/10.1214/aos/1013203451.",
    "7. Chen, T.; Guestrin, C. XGBoost: A Scalable Tree Boosting System. Proc. KDD 2016, 785-794. https://doi.org/10.1145/2939672.2939785.",
    "8. Rasmussen, C. E.; Williams, C. K. I. Gaussian Processes for Machine Learning; MIT Press: Cambridge, MA, 2006.",
    "9. Lundberg, S. M.; Lee, S.-I. A Unified Approach to Interpreting Model Predictions. Adv. Neural Inf. Process. Syst. 2017, 30, 4765-4774.",
    "10. Efron, B.; Tibshirani, R. J. An Introduction to the Bootstrap; Chapman & Hall/CRC: New York, 1993.",
]
for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.20)
    p.paragraph_format.first_line_indent = Inches(-0.20)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(ref), size=8.6)

doc.core_properties.title = (
    "Supplementary Information - Reproducible Physics-Constrained "
    "LLM-Assisted Symbolic Regression for Hybrid Perovskite Band Gaps"
)
doc.core_properties.subject = "Supplementary evidence and reproducibility package"
doc.core_properties.author = ""
doc.core_properties.keywords = (
    "hybrid perovskites, symbolic regression, LLM, supplementary information, reproducibility"
)
doc.core_properties.comments = "Generated from audited local project outputs."

doc.save(OUTPUT)
print(OUTPUT)
